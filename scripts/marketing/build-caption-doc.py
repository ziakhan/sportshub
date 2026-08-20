"""
Builds the shareable reference sheet for the Instagram launch: every post as a
row, thumbnail beside its caption, so the team can scan and edit copy.

    python3 scripts/marketing/build-caption-doc.py [outPath] [shotsDir]

NOT a deck. An earlier version gave each post a full page with a large image,
which is the wrong artefact: the owner already holds the high-res exports for
posting, so the document only has to say which caption goes with which card.
Thumbnails are downscaled to ~460px before embedding, which keeps the file
small enough to sit in a shared doc without fighting it.

Single source of truth is docs/marketing/instagram-launch-captions-2026-08.md,
parsed here rather than retyped, so the document cannot drift from the copy
that ships. Images come from a render of the creatives, so run
render-creatives.mjs first if any card has changed.

.docx rather than PDF: Google Docs imports it with images intact, so captions
stay editable in the shared copy.
"""
import io
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image

REPO = Path(__file__).resolve().parents[2]
DOC = REPO / "docs/marketing/instagram-launch-captions-2026-08.md"
SHOTS = Path(sys.argv[2]) if len(sys.argv) > 2 else Path(
    "/private/tmp/claude-501/-Users-ziakhan-zia-personal-sportshub"
    "/3a61122f-9c3c-432b-b6eb-701864840679/scratchpad/creatives"
)
OUT = Path(sys.argv[1]) if len(sys.argv) > 1 else REPO / "instagram-launch-posts.docx"

BASE_TAGS = "#youthbasketball #ontariobasketball #gtabasketball #basketballparents"
EXTRA_TAGS = {"7": "#basketballleague", "9": "#basketballtrainer", "10": "#basketballtryouts"}

NAVY = RGBColor(0x0B, 0x16, 0x28)
GREY = RGBColor(0x74, 0x74, 0x86)
ORANGE = RGBColor(0xF2, 0x4E, 0x1E)

THUMB_PX = 460          # embedded pixel width; display width is set separately
THUMB_IN = Inches(1.55)  # on the page


def parse():
    raw = DOC.read_text()
    posts = []
    for block in re.split(r"^## ", raw, flags=re.M)[1:]:
        head = block.split("\n", 1)[0]
        m = re.match(r"^(\d+)\s+·\s+(.+)$", head)
        if not m:
            continue
        slug = re.search(r"`([a-z0-9-]+)`", block)
        if not slug:
            continue
        ask = re.search(r"\*\*Ask: (\w+)\*\*", block)
        quoted = [l[2:].rstrip() if l.startswith("> ") else "" for l in block.split("\n") if l.startswith(">")]
        paras = [" ".join(p.split()) for p in "\n".join(quoted).split("\n\n") if p.strip()]
        posts.append({
            "no": m.group(1), "title": m.group(2).strip(), "slug": slug.group(1),
            "ask": (ask.group(1) if ask else "reply"), "paras": paras,
        })
    return sorted(posts, key=lambda p: int(p["no"]))


def thumb(path: Path) -> io.BytesIO:
    """Downscale before embedding. python-docx stores the bytes you give it,
    so setting a small display width on a 1080px PNG still ships 1080px."""
    im = Image.open(path).convert("RGB")
    im.thumbnail((THUMB_PX, THUMB_PX * 4), Image.LANCZOS)
    buf = io.BytesIO()
    im.save(buf, format="JPEG", quality=82, optimize=True)
    buf.seek(0)
    return buf


def main():
    posts = parse()
    d = Document()
    for s in d.sections:
        s.top_margin = s.bottom_margin = Inches(0.6)
        s.left_margin = s.right_margin = Inches(0.7)

    h = d.add_heading("SportsHub One — Instagram launch", level=0)
    h.runs[0].font.color.rgb = NAVY
    p = d.add_paragraph()
    r = p.add_run(
        f"{len(posts)} posts in order. Thumbnails are for reference only; post from the "
        "full-size exports. One ask per post: “reply” asks a question to earn comments, "
        "“convert” asks for the click."
    )
    r.font.size = Pt(9.5)
    r.font.color.rgb = GREY

    table = d.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    missing = []

    for post in posts:
        row = table.add_row()
        left, right = row.cells
        left.width = Inches(1.7)
        right.width = Inches(5.4)

        img = SHOTS / f"{post['slug']}-portrait.png"
        cell_p = left.paragraphs[0]
        cell_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if img.exists():
            cell_p.add_run().add_picture(thumb(img), width=THUMB_IN)
        else:
            missing.append(post["slug"])
            cell_p.add_run(f"[no render: {post['slug']}]").font.size = Pt(8)

        head_p = right.paragraphs[0]
        run = head_p.add_run(f"Post {post['no']} — {post['title']}")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = NAVY

        meta = right.add_paragraph()
        r = meta.add_run(f"{post['slug']}  ·  ask: {post['ask']}")
        r.font.size = Pt(8)
        r.font.color.rgb = GREY

        for i, para in enumerate(post["paras"]):
            el = right.add_paragraph()
            run = el.add_run(para)
            run.font.size = Pt(10.5)
            if i == len(post["paras"]) - 1 and len(post["paras"]) > 1:
                run.bold = True
                run.font.color.rgb = ORANGE

        tags = right.add_paragraph()
        r = tags.add_run(BASE_TAGS + (" " + EXTRA_TAGS[post["no"]] if post["no"] in EXTRA_TAGS else ""))
        r.font.size = Pt(8.5)
        r.font.color.rgb = GREY

    OUT.parent.mkdir(parents=True, exist_ok=True)
    d.save(OUT)
    print(f"wrote {OUT}  ({len(posts)} posts, {OUT.stat().st_size/1024:.0f} KB)")
    if missing:
        print("MISSING RENDERS:", ", ".join(missing))


if __name__ == "__main__":
    main()
